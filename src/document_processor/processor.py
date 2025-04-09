import io
import os
from typing import Union

import PyPDF2
import docx

class DocumentProcessor:
    def __init__(self):
        self.supported_types = {
            "application/pdf": self._process_pdf,
            "text/plain": self._process_txt,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._process_docx
        }
    
    def process_document(self, file: Union[io.BytesIO, str]) -> str:
        """Process uploaded document and return text content"""
        if isinstance(file, str):
            with open(file, 'rb') as f:
                file = io.BytesIO(f.read())
        
        file_type = self._get_file_type(file)
        if file_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        try:
            return self.supported_types[file_type](file)
        except Exception as e:
            # If processing fails, try as text
            file.seek(0)
            try:
                return self._process_txt(file)
            except:
                raise ValueError(f"Failed to process file: {str(e)}")
    
    def _get_file_type(self, file: io.BytesIO) -> str:
        """Get file type from file object"""
        # Check if file has a name attribute (from file upload)
        if hasattr(file, 'name'):
            if file.name.endswith('.pdf'):
                return "application/pdf"
            elif file.name.endswith('.txt'):
                return "text/plain"
            elif file.name.endswith('.docx'):
                return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # If no name attribute, try to determine type from content
        try:
            # Try to read as PDF
            PyPDF2.PdfReader(file)
            file.seek(0)  # Reset file pointer
            return "application/pdf"
        except:
            try:
                # Try to read as DOCX
                docx.Document(file)
                file.seek(0)  # Reset file pointer
                return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            except:
                # Assume it's a text file
                file.seek(0)  # Reset file pointer
                return "text/plain"
    
    def _process_pdf(self, file: io.BytesIO) -> str:
        """Process PDF file and return text content"""
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    def _process_txt(self, file: io.BytesIO) -> str:
        """Process text file and return content"""
        content = file.read().decode('utf-8')
        return content if content.endswith('\n') else content + '\n'
    
    def _process_docx(self, file: io.BytesIO) -> str:
        """Process DOCX file and return text content"""
        try:
            doc = docx.Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except:
            # If DOCX processing fails, try as text
            file.seek(0)
            return self._process_txt(file) 